"""
Generates realistic .docx (and then .pdf via LibreOffice) resume files
from plain-text resume samples, so the extraction engine has real files
to be tested against instead of just .txt.
"""

import subprocess
from pathlib import Path
from docx import Document
from docx.shared import Pt

SOURCE_DIR = Path("/home/claude/day4/samples/resumes")
DOCX_DIR = Path("/home/claude/day5/data/raw_resumes/docx")
PDF_DIR = Path("/home/claude/day5/data/raw_resumes/pdf")

DOCX_DIR.mkdir(parents=True, exist_ok=True)
PDF_DIR.mkdir(parents=True, exist_ok=True)

SECTION_HEADERS = {"Summary:", "Skills:", "Experience:", "Education:", "Certifications:"}


def txt_to_docx(txt_path: Path, docx_path: Path) -> None:
    lines = txt_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    # First non-empty line is the candidate name -> Heading 1
    name_written = False

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            doc.add_paragraph("")
            continue

        if not name_written and line.startswith("Name:"):
            name = line.replace("Name:", "").strip()
            doc.add_heading(name, level=1)
            name_written = True
            continue

        if line in SECTION_HEADERS:
            doc.add_heading(line.rstrip(":"), level=2)
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
            continue

        p = doc.add_paragraph(line)

    doc.save(docx_path)


def main():
    txt_files = sorted(SOURCE_DIR.glob("resume_*.txt"))
    print(f"Found {len(txt_files)} source resumes")

    for txt_path in txt_files:
        base_name = txt_path.stem
        docx_path = DOCX_DIR / f"{base_name}.docx"
        txt_to_docx(txt_path, docx_path)
        print(f"  created {docx_path.name}")

    # Convert every docx to pdf in one batch via LibreOffice headless
    print("Converting DOCX -> PDF via LibreOffice...")
    result = subprocess.run(
        [
            "python", "/mnt/skills/public/docx/scripts/office/soffice.py",
            "--headless", "--convert-to", "pdf",
            "--outdir", str(PDF_DIR),
        ] + [str(p) for p in DOCX_DIR.glob("*.docx")],
        capture_output=True, text=True,
    )
    print(result.stdout[-1500:])
    if result.returncode != 0:
        print("STDERR:", result.stderr[-1500:])


if __name__ == "__main__":
    main()

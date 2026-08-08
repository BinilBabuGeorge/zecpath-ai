"""
Resume Text Extraction Engine (Day 5)

Reads a resume file (.pdf or .docx), extracts raw text, and normalizes it
into a clean, consistent format ready for downstream AI parsing
(the ats_engine / parsers.resume_parser layer built in Day 3-4).

Public API:
    extract_resume(file_path) -> ExtractionResult
"""

from __future__ import annotations

import re
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path
from typing import List

import pdfplumber
from docx import Document

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}

# Section headings we expect to see on a resume (used to normalize casing/format)
KNOWN_SECTIONS = ["Summary", "Skills", "Experience", "Education", "Certifications"]

# Characters commonly used as bullet markers across different resume exports
BULLET_CHARS = ["•", "◦", "▪", "‣", "·", "●", "*", "-", "–"]


@dataclass
class ExtractionResult:
    file_path: str
    file_type: str
    raw_text: str
    cleaned_text: str
    line_count: int
    detected_sections: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Raw extraction (format-specific readers)
# ---------------------------------------------------------------------------

def _read_pdf(path: Path) -> str:
    """Extract raw text from a PDF, page by page. Handles multi-column pages
    reasonably well because pdfplumber reads in visual (left-to-right,
    top-to-bottom) order per text block."""
    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    return "\n".join(text_parts)


def _read_docx(path: Path) -> str:
    """Extract raw text from a DOCX, including paragraphs and table cells
    (some resumes lay out skills/experience inside tables)."""
    doc = Document(path)
    text_parts = [p.text for p in doc.paragraphs]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)

    return "\n".join(text_parts)


def extract_raw_text(file_path: str) -> str:
    """Dispatch to the correct reader based on file extension."""
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type '{ext}'. Supported: {SUPPORTED_EXTENSIONS}")

    if ext == ".pdf":
        return _read_pdf(path)
    if ext == ".docx":
        return _read_docx(path)
    raise AssertionError("unreachable")  # SUPPORTED_EXTENSIONS guards this


# ---------------------------------------------------------------------------
# Cleaning & normalization
# ---------------------------------------------------------------------------

def _normalize_unicode(text: str) -> str:
    """Fold curly quotes, non-breaking spaces, etc. into plain ASCII-friendly forms."""
    text = unicodedata.normalize("NFKC", text)
    replacements = {
        "\u2018": "'", "\u2019": "'",   # curly single quotes
        "\u201c": '"', "\u201d": '"',   # curly double quotes
        "\u00a0": " ",                   # non-breaking space
        "\u2013": "-", "\u2014": "-",   # en/em dash
    }
    for bad, good in replacements.items():
        text = text.replace(bad, good)
    return text


def _normalize_bullets(text: str) -> str:
    """Convert every bullet-marker style into a single consistent '- ' prefix."""
    lines = text.split("\n")
    normalized = []
    for line in lines:
        stripped = line.strip()
        for bullet in BULLET_CHARS:
            if stripped.startswith(bullet + " ") or stripped == bullet:
                stripped = "- " + stripped[len(bullet):].strip()
                break
        normalized.append(stripped)
    return "\n".join(normalized)


def _normalize_section_headings(text: str) -> str:
    """Standardize known section headings to 'Title Case:' regardless of how
    they were capitalized or punctuated in the source file."""
    lines = text.split("\n")
    normalized = []
    for line in lines:
        stripped = line.strip().rstrip(":")
        matched = None
        for section in KNOWN_SECTIONS:
            if stripped.lower() == section.lower():
                matched = section
                break
        normalized.append(f"{matched}:" if matched else line)
    return "\n".join(normalized)


def _remove_noise(text: str) -> str:
    """Strip control characters, collapse repeated blank lines/spaces,
    and drop lines that are just leftover formatting artifacts."""
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)  # control chars
    text = re.sub(r"[ \t]+", " ", text)                        # repeated spaces/tabs
    text = re.sub(r"\n{3,}", "\n\n", text)                     # 3+ blank lines -> 1
    lines = [ln for ln in text.split("\n") if ln.strip() not in {"", "|", "_", "-"}]
    return "\n".join(line.strip() for line in lines)


def clean_text(raw_text: str) -> str:
    """Full cleaning/normalization pipeline applied to raw extracted text."""
    text = _normalize_unicode(raw_text)
    text = _normalize_bullets(text)
    text = _normalize_section_headings(text)
    text = _remove_noise(text)
    return text.strip()


def detect_sections(cleaned_text: str) -> List[str]:
    """Return which known resume sections were found in the cleaned text."""
    found = []
    for section in KNOWN_SECTIONS:
        if re.search(rf"^{section}:", cleaned_text, re.MULTILINE):
            found.append(section)
    return found


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def extract_resume(file_path: str) -> ExtractionResult:
    """Read a resume file end-to-end: extract raw text, clean it, and report
    which expected sections were detected."""
    path = Path(file_path)
    warnings: List[str] = []

    raw_text = extract_raw_text(str(path))
    if not raw_text.strip():
        warnings.append("No extractable text found (file may be a scanned image).")

    cleaned = clean_text(raw_text)
    sections = detect_sections(cleaned)

    missing = [s for s in KNOWN_SECTIONS if s not in sections]
    if missing:
        warnings.append(f"Missing expected sections: {', '.join(missing)}")

    return ExtractionResult(
        file_path=str(path),
        file_type=path.suffix.lower().lstrip("."),
        raw_text=raw_text,
        cleaned_text=cleaned,
        line_count=len(cleaned.splitlines()),
        detected_sections=sections,
        warnings=warnings,
    )


if __name__ == "__main__":
    import sys
    result = extract_resume(sys.argv[1])
    print(f"File: {result.file_path}")
    print(f"Sections found: {result.detected_sections}")
    print(f"Warnings: {result.warnings}")
    print("---- Cleaned text ----")
    print(result.cleaned_text)
